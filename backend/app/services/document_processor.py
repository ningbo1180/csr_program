"""
Document Processing Service
"""
import os
import logging
from pathlib import Path
from typing import List, Dict, Any
import mimetypes
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Base document processor"""
    
    SUPPORTED_FORMATS = [".pdf", ".docx", ".xlsx"]
    MAX_FILE_SIZE = 52428800  # 50MB
    
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def validate_file(self, filename: str, file_size: int) -> tuple[bool, str]:
        """Validate file before processing"""
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            return False, f"Unsupported file format: {file_ext}"
        
        # Check file size
        if file_size > self.MAX_FILE_SIZE:
            return False, f"File size exceeds limit: {file_size} > {self.MAX_FILE_SIZE}"
        
        return True, "OK"
    
    def save_file(self, project_id: str, filename: str, file_content: bytes) -> str:
        """Save uploaded file to disk"""
        project_dir = self.upload_dir / project_id
        project_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        file_path = project_dir / f"{timestamp}_{filename}"
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        logger.info(f"File saved: {file_path}")
        return str(file_path)
    
    async def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file and extract content"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            chapters = []
            text_content = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                text_content.append(text)
                
                # Simple chapter detection by looking for numbered sections
                lines = text.split("\n")
                for line in lines:
                    if line.strip() and (line[0].isdigit() or line.startswith("#")):
                        chapters.append({
                            "page": page_num + 1,
                            "text": line.strip()[:100]
                        })
            
            doc.close()
            
            return {
                "success": True,
                "chapters": chapters[:20],  # Limit to 20 chapters
                "total_pages": len(doc),
                "content_preview": "\n".join(text_content[:500])
            }
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX file and extract content"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            chapters = []
            for para in doc.paragraphs:
                # Extract paragraphs that look like chapter titles
                if para.style.name.startswith("Heading") or (
                    para.text.strip() and 
                    para.text[0].isdigit() and 
                    len(chapters) < 20
                ):
                    chapters.append({
                        "text": para.text.strip()[:100],
                        "style": para.style.name
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data[:5])  # Limit to 5 rows per table
            
            return {
                "success": True,
                "chapters": chapters,
                "tables_count": len(tables),
                "tables_sample": tables[:3]
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Process XLSX file and extract table data"""
        try:
            import pandas as pd
            
            excel_file = pd.ExcelFile(file_path)
            sheets = excel_file.sheet_names
            
            tables = []
            for sheet in sheets[:5]:  # Limit to 5 sheets
                df = pd.read_excel(file_path, sheet_name=sheet)
                tables.append({
                    "sheet_name": sheet,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns)
                })
            
            return {
                "success": True,
                "sheets": sheets,
                "tables": tables,
                "total_sheets": len(sheets)
            }
        except Exception as e:
            logger.error(f"Error processing XLSX: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process document based on file type"""
        if file_type == "pdf":
            return await self.process_pdf(file_path)
        elif file_type == "docx":
            return await self.process_docx(file_path)
        elif file_type == "xlsx":
            return await self.process_xlsx(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}"
            }
