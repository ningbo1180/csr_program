"""
Export API routes - Generate CSR Word document
"""
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import datetime
import logging
import io

from app.database import get_db
from app.models.models import Chapter, Project

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/export", tags=["export"])


def _build_chapter_tree(db: Session, project_id: str) -> list:
    """Build flat ordered list of chapters from tree"""
    root_chapters = db.query(Chapter).filter(
        Chapter.project_id == project_id,
        Chapter.parent_id == None
    ).order_by(Chapter.order, Chapter.number).all()
    
    result = []
    
    def traverse(chapter, level=0):
        result.append({
            "id": chapter.id,
            "title": chapter.title,
            "number": chapter.number,
            "content": chapter.content or "",
            "level": level,
        })
        children = db.query(Chapter).filter(
            Chapter.parent_id == chapter.id
        ).order_by(Chapter.order, Chapter.number).all()
        for child in children:
            traverse(child, level + 1)
    
    for ch in root_chapters:
        traverse(ch, 0)
    
    return result


def _strip_html_tags(html: str) -> str:
    """Simple HTML tag stripper for plain text fallback"""
    import re
    text = re.sub(r'<br\s*/?>', '\n', html)
    text = re.sub(r'<p[^>]*>', '\n', text)
    text = re.sub(r'</p>', '', text)
    text = re.sub(r'<div[^>]*>', '\n', text)
    text = re.sub(r'</div>', '', text)
    text = re.sub(r'<li[^>]*>', '\n• ', text)
    text = re.sub(r'</li>', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


@router.get("/{project_id}")
async def export_csr_docx(
    project_id: str,
    db: Session = Depends(get_db)
):
    """Export project CSR as Word document (.docx)"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Clinical Study Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(project.name)
    run2.font.size = Pt(16)
    
    # Optional: add study info if available in project metadata
    if project.description:
        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc.add_run(project.description).font.size = Pt(12)
    
    doc.add_paragraph()  # spacing
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").font.size = Pt(10)
    meta.add_run(f"Language: {project.language or 'zh-CN'}\n").font.size = Pt(10)
    status_val = project.status.value if hasattr(project.status, 'value') else str(project.status) if project.status else 'draft'
    meta.add_run(f"Status: {status_val}").font.size = Pt(10)
    
    doc.add_page_break()
    
    # Table of Contents header
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    chapters = _build_chapter_tree(db, project_id)
    
    # Build TOC
    for ch in chapters:
        indent = "    " * ch["level"]
        toc_line = doc.add_paragraph()
        toc_line.add_run(f"{indent}{ch['number']}  {ch['title']}")
    
    doc.add_page_break()
    
    # Chapter contents
    for ch in chapters:
        level = min(ch["level"] + 1, 3)
        heading = doc.add_heading(f"{ch['number']}  {ch['title']}", level=level)
        
        content = ch["content"]
        if content:
            # Try to preserve some HTML structure
            # Split by common block elements
            blocks = content.replace('</p>', '\n').replace('</div>', '\n').replace('<br>', '\n').replace('<br/>', '\n').split('\n')
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                # Check if it's a table
                if '<table' in block.lower():
                    # Extract table rows
                    import re
                    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', block, re.DOTALL)
                    if rows:
                        # Count columns from first row
                        first_row_cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', rows[0], re.DOTALL)
                        if first_row_cells:
                            table = doc.add_table(rows=1, cols=len(first_row_cells))
                            table.style = 'Table Grid'
                            # Header row
                            for i, cell_html in enumerate(first_row_cells):
                                cell_text = _strip_html_tags(cell_html)
                                table.rows[0].cells[i].text = cell_text
                            # Data rows
                            for row_html in rows[1:]:
                                cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL)
                                if cells:
                                    row = table.add_row()
                                    for i, cell_html in enumerate(cells):
                                        if i < len(first_row_cells):
                                            row.cells[i].text = _strip_html_tags(cell_html)
                            doc.add_paragraph()
                else:
                    # Regular paragraph
                    text = _strip_html_tags(block)
                    if text:
                        p = doc.add_paragraph(text)
                        p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph("[Content pending generation]")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            p.runs[0].italic = True
    
    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"CSR_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
